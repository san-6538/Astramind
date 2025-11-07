import os
import io
import fitz  # PyMuPDF
import pdfplumber
from datetime import datetime
from docx import Document
from typing import List
from PIL import Image

from services.ocr import multi_stage_ocr

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ============================================================
# ✅ File Saving
# ============================================================
def save_upload(file) -> str:
    """
    Save uploaded file locally in the /uploads directory.
    """
    path = os.path.join(UPLOAD_DIR, file.filename)
    with open(path, "wb") as f:
        f.write(file.file.read())
    print(f"📂 Saved upload: {file.filename}")
    return path


# ============================================================
# ✅ Smarter Text Chunking
# ============================================================
def chunk_text(text: str, max_tokens: int = 700, overlap: int = 80) -> List[str]:
    """
    Adaptive chunking for large documents.
    Splits at sentence boundaries when possible for coherence.
    """
    import re

    if not text.strip():
        return []

    # Normalize spaces and split sentences
    sentences = re.split(r'(?<=[.!?]) +', text)
    chunks, current_chunk = [], []

    token_count = 0
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        # Roughly estimate token count (4 chars ≈ 1 token)
        tokens = len(sentence) // 4
        if token_count + tokens > max_tokens:
            # Create chunk when limit exceeded
            chunks.append(" ".join(current_chunk).strip())
            # Keep small overlap
            overlap_text = current_chunk[-2:] if len(current_chunk) >= 2 else current_chunk[-1:]
            current_chunk = overlap_text.copy()
            token_count = sum(len(s) // 4 for s in current_chunk)

        current_chunk.append(sentence)
        token_count += tokens

    if current_chunk:
        chunks.append(" ".join(current_chunk).strip())

    print(f"✂️ Chunked text into {len(chunks)} parts (~{max_tokens} tokens each).")
    return chunks


# ============================================================
# ✅ Enhanced PDF Extractor (Multi-column + OCR + Tables)
# ============================================================
def extract_text_from_pdf(pdf_path: str, max_pages: int = 25) -> List[str]:
    """
    Extracts clean, chunked text from PDFs.
    Handles tables, multi-columns, and OCR fallback for scanned pages.
    Efficient enough for 20–25 page documents.
    """
    chunks = []
    try:
        print(f"📘 Extracting text from PDF: {pdf_path}")
        with pdfplumber.open(pdf_path) as pdf:
            page_count = min(len(pdf.pages), max_pages)
            for page_num, page in enumerate(pdf.pages[:page_count], start=1):
                try:
                    text = page.extract_text(layout=True) or ""
                    tables = page.extract_tables() or []

                    # Add table data as Markdown
                    for table in tables:
                        table_text = "\n".join(
                            [" | ".join(str(cell).strip() for cell in row if cell) for row in table]
                        )
                        if table_text.strip():
                            text += "\n" + table_text

                    # OCR fallback for scanned pages
                    if not text.strip():
                        print(f"⚠️ Page {page_num} is scanned → running OCR...")
                        image = page.to_image(resolution=300).original
                        temp_path = f"{pdf_path}_page_{page_num}.png"
                        image.save(temp_path)
                        ocr_result = multi_stage_ocr(temp_path, "image")
                        text = ocr_result.get("text", "")
                        os.remove(temp_path)

                    # Chunk and add to list
                    if text.strip():
                        chunks.extend(chunk_text(text))

                except Exception as e:
                    print(f"⚠️ Error processing PDF page {page_num}: {e}")

    except Exception as e:
        print(f"❌ PDF extraction error: {e}")

    print(f"✅ Extracted {len(chunks)} chunks from {min(len(chunks), max_pages)} pages.")
    return chunks


# ============================================================
# ✅ DOCX Extractor (Paragraphs + Tables + Embedded Images)
# ============================================================
def extract_text_from_docx(file) -> List[str]:
    """
    Extracts text, tables, and OCR from embedded images in DOCX.
    Supports large, multi-section Word documents.
    """
    text_chunks = []
    try:
        doc = Document(io.BytesIO(file.file.read()))
        print("📄 Extracting text from DOCX...")

        # 1️⃣ Paragraphs
        para_buffer = []
        for para in doc.paragraphs:
            if para.text.strip():
                para_buffer.append(para.text.strip())
            if len(para_buffer) > 5:
                text_chunks.extend(chunk_text(" ".join(para_buffer)))
                para_buffer = []
        if para_buffer:
            text_chunks.extend(chunk_text(" ".join(para_buffer)))

        # 2️⃣ Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_chunks.append(f"TABLE_ROW: {row_text}")

        # 3️⃣ Embedded images (OCR)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    temp_path = os.path.join(UPLOAD_DIR, "temp_docx_img.png")
                    with open(temp_path, "wb") as img_file:
                        img_file.write(image_data)

                    ocr_result = multi_stage_ocr(temp_path, "image")
                    ocr_text = ocr_result.get("text", "")
                    if ocr_text.strip():
                        text_chunks.append(f"OCR_IMAGE: {ocr_text.strip()}")

                    os.remove(temp_path)
                except Exception as e:
                    print(f"⚠️ OCR image extraction failed: {e}")

        print(f"✅ Extracted {len(text_chunks)} chunks from DOCX.")
    except Exception as e:
        print(f"❌ DOCX extraction error: {e}")

    return text_chunks


# ============================================================
# ✅ Image Extractor (Tesseract → EasyOCR → Gemini Vision)
# ============================================================
def extract_text_from_image(file_or_path):
    """
    Extracts text from an image or in-memory file using hybrid OCR.
    Handles large or multi-page scanned images.
    """
    try:
        if isinstance(file_or_path, (str, os.PathLike)):
            result = multi_stage_ocr(file_or_path, "image")
        else:
            temp_path = os.path.join(UPLOAD_DIR, "temp_upload_image.png")
            with open(temp_path, "wb") as f:
                f.write(file_or_path.read())
            result = multi_stage_ocr(temp_path, "image")
            os.remove(temp_path)
        return result.get("text", "")
    except Exception as e:
        print(f"❌ Image OCR error: {e}")
        return ""
